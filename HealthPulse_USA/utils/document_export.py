"""
Document Generation Utilities
Handles export to PDF, DOCX, and TXT formats
"""

import io
import re
import os
from typing import Optional
from datetime import datetime

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def sanitize_text_for_pdf(text: str) -> str:
    """
    Replace Unicode characters that aren't supported by standard PDF fonts
    with ASCII equivalents
    """
    replacements = {
        '•': '-',      # Bullet point
        '–': '-',      # En dash
        '—': '-',      # Em dash
        ''': "'",      # Left single quote
        ''': "'",      # Right single quote
        '"': '"',      # Left double quote
        '"': '"',      # Right double quote
        '…': '...',    # Ellipsis
        '©': '(c)',    # Copyright
        '®': '(R)',    # Registered
        '™': '(TM)',   # Trademark
        '°': ' deg',   # Degree
        '±': '+/-',    # Plus-minus
        '×': 'x',      # Multiplication
        '÷': '/',      # Division
        '≤': '<=',     # Less than or equal
        '≥': '>=',     # Greater than or equal
        '≠': '!=',     # Not equal
        '→': '->',     # Right arrow
        '←': '<-',     # Left arrow
        '✓': '[x]',    # Checkmark
        '✗': '[ ]',    # X mark
        '★': '*',      # Star
        '✅': '[PASS]', # Green checkmark
        '❌': '[FAIL]', # Red X
    }
    
    for unicode_char, ascii_char in replacements.items():
        text = text.replace(unicode_char, ascii_char)
    
    # Remove any remaining non-ASCII characters
    text = text.encode('ascii', 'ignore').decode('ascii')
    
    return text


class ArticleExporter:
    """Exports articles to various formats"""
    
    def __init__(self, title: str, content: str, meta_description: str = "", 
                 seo_score: float = 0.0, keywords: list = None):
        """
        Initialize the exporter with article data
        
        Args:
            title: Article title
            content: Full article content (markdown format)
            meta_description: SEO meta description
            seo_score: SEO validation score
            keywords: List of keywords
        """
        self.title = title
        self.content = content
        self.meta_description = meta_description
        self.seo_score = seo_score
        self.keywords = keywords or []
        self.generation_date = datetime.now().strftime("%B %d, %Y")
    
    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting for plain text"""
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        # Remove bold/italic
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Remove links
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        # Clean up extra whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def _parse_markdown_sections(self, content: str) -> list:
        """Parse markdown content into sections with headers"""
        sections = []
        current_section = {"level": 0, "title": "", "content": ""}
        
        for line in content.split('\n'):
            if line.startswith('# '):
                if current_section["content"] or current_section["title"]:
                    sections.append(current_section)
                current_section = {"level": 1, "title": line[2:].strip(), "content": ""}
            elif line.startswith('## '):
                if current_section["content"] or current_section["title"]:
                    sections.append(current_section)
                current_section = {"level": 2, "title": line[3:].strip(), "content": ""}
            elif line.startswith('### '):
                if current_section["content"] or current_section["title"]:
                    sections.append(current_section)
                current_section = {"level": 3, "title": line[4:].strip(), "content": ""}
            else:
                current_section["content"] += line + "\n"
        
        if current_section["content"] or current_section["title"]:
            sections.append(current_section)
        
        return sections
    
    def export_to_txt(self) -> bytes:
        """Export article to plain text format"""
        output = io.StringIO()
        
        # Header
        output.write("=" * 70 + "\n")
        output.write(f"  {self.title}\n")
        output.write("=" * 70 + "\n\n")
        
        # Meta information
        output.write(f"Generated: {self.generation_date}\n")
        output.write(f"SEO Score: {self.seo_score}%\n")
        if self.keywords:
            output.write(f"Keywords: {', '.join(self.keywords)}\n")
        output.write("\n" + "-" * 70 + "\n\n")
        
        # Meta description
        if self.meta_description:
            output.write("META DESCRIPTION:\n")
            output.write(self.meta_description + "\n\n")
            output.write("-" * 70 + "\n\n")
        
        # Content
        output.write(self._clean_markdown(self.content))
        
        # Footer
        output.write("\n\n" + "=" * 70 + "\n")
        output.write("Generated by HealthPulse USA - Multi-Agent Healthcare Content System\n")
        output.write("=" * 70 + "\n")
        
        return output.getvalue().encode('utf-8')
    
    def export_to_docx(self) -> Optional[bytes]:
        """Export article to DOCX format"""
        if not DOCX_AVAILABLE:
            return None
        
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Title style
        title_style = styles['Title']
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add title
        title_para = doc.add_paragraph(self.title, style='Title')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add meta information
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {self.generation_date} | SEO Score: {self.seo_score}%")
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        doc.add_paragraph("_" * 60)
        
        # Add meta description
        if self.meta_description:
            doc.add_paragraph()
            meta_heading = doc.add_paragraph()
            meta_heading_run = meta_heading.add_run("Meta Description")
            meta_heading_run.bold = True
            meta_heading_run.font.size = Pt(11)
            
            meta_content = doc.add_paragraph(self.meta_description)
            meta_content.paragraph_format.left_indent = Inches(0.25)
            for run in meta_content.runs:
                run.font.italic = True
                run.font.size = Pt(10)
            
            doc.add_paragraph()
        
        # Add keywords
        if self.keywords:
            keywords_para = doc.add_paragraph()
            keywords_run = keywords_para.add_run(f"Keywords: {', '.join(self.keywords)}")
            keywords_run.font.size = Pt(10)
            keywords_run.font.color.rgb = RGBColor(0, 102, 153)
            doc.add_paragraph()
        
        # Parse and add content sections
        sections = self._parse_markdown_sections(self.content)
        
        for section in sections:
            if section["title"]:
                if section["level"] == 1:
                    heading = doc.add_heading(section["title"], level=1)
                    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                elif section["level"] == 2:
                    heading = doc.add_heading(section["title"], level=2)
                    heading.runs[0].font.color.rgb = RGBColor(0, 76, 153)
                elif section["level"] == 3:
                    heading = doc.add_heading(section["title"], level=3)
                    heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            
            # Process content
            content = section["content"].strip()
            if content:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Check for bullet points
                        if para_text.startswith('- ') or para_text.startswith('* '):
                            lines = para_text.split('\n')
                            for line in lines:
                                if line.strip().startswith(('- ', '* ')):
                                    bullet_text = line.strip()[2:]
                                    para = doc.add_paragraph(bullet_text, style='List Bullet')
                                elif line.strip():
                                    doc.add_paragraph(line.strip())
                        # Check for numbered lists
                        elif para_text[0].isdigit() and para_text[1:3] in ['. ', ') ']:
                            lines = para_text.split('\n')
                            for line in lines:
                                if line.strip() and line.strip()[0].isdigit():
                                    # Extract text after number
                                    list_text = re.sub(r'^\d+[\.\)]\s*', '', line.strip())
                                    para = doc.add_paragraph(list_text, style='List Number')
                                elif line.strip():
                                    doc.add_paragraph(line.strip())
                        else:
                            # Regular paragraph - handle bold text
                            para = doc.add_paragraph()
                            # Split on bold markers
                            parts = re.split(r'(\*\*.*?\*\*)', para_text.replace('\n', ' '))
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = para.add_run(part[2:-2])
                                    run.bold = True
                                else:
                                    para.add_run(part)
        
        # Add footer
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)
        footer = doc.add_paragraph()
        footer_run = footer.add_run("Generated by HealthPulse USA - Multi-Agent Healthcare Content System")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_pdf(self) -> Optional[bytes]:
        """Export article to PDF format"""
        if not PDF_AVAILABLE:
            return None
        
        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=20)
            pdf.add_page()
            
            # Use only ASCII-safe content
            safe_title = sanitize_text_for_pdf(self.title)[:80]  # Limit title length
            safe_meta = sanitize_text_for_pdf(self.meta_description)[:200]
            safe_content = sanitize_text_for_pdf(self.content)
            
            # Title - centered, limited width
            pdf.set_font('Helvetica', 'B', 18)
            pdf.set_text_color(0, 51, 102)
            
            # Calculate title width and wrap if needed
            title_width = pdf.get_string_width(safe_title)
            if title_width > 170:
                safe_title = safe_title[:60] + "..."
            
            pdf.cell(0, 12, safe_title, ln=True, align='C')
            pdf.ln(5)
            
            # Meta info line
            pdf.set_font('Helvetica', 'I', 9)
            pdf.set_text_color(128, 128, 128)
            meta_line = f"Generated: {self.generation_date} | SEO Score: {self.seo_score}%"
            pdf.cell(0, 6, meta_line, ln=True, align='C')
            pdf.ln(8)
            
            # Horizontal line
            pdf.set_draw_color(200, 200, 200)
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.ln(8)
            
            # Meta description
            if safe_meta:
                pdf.set_font('Helvetica', 'B', 10)
                pdf.set_text_color(0, 0, 0)
                pdf.cell(0, 6, "Meta Description:", ln=True)
                pdf.set_font('Helvetica', 'I', 9)
                pdf.set_text_color(80, 80, 80)
                pdf.multi_cell(0, 5, safe_meta)
                pdf.ln(5)
            
            # Main content
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(0, 0, 0)
            
            # Process content line by line
            lines = safe_content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    pdf.ln(3)
                    continue
                
                # Handle headers
                if line.startswith('# '):
                    pdf.ln(5)
                    pdf.set_font('Helvetica', 'B', 14)
                    pdf.set_text_color(0, 51, 102)
                    header_text = line[2:].strip()[:70]
                    pdf.multi_cell(0, 7, header_text)
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(2)
                elif line.startswith('## '):
                    pdf.ln(4)
                    pdf.set_font('Helvetica', 'B', 12)
                    pdf.set_text_color(0, 76, 153)
                    header_text = line[3:].strip()[:70]
                    pdf.multi_cell(0, 6, header_text)
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(2)
                elif line.startswith('### '):
                    pdf.ln(3)
                    pdf.set_font('Helvetica', 'B', 11)
                    pdf.set_text_color(51, 102, 153)
                    header_text = line[4:].strip()[:70]
                    pdf.multi_cell(0, 5, header_text)
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(1)
                elif line.startswith('- ') or line.startswith('* '):
                    bullet_text = "  - " + line[2:].strip()
                    pdf.multi_cell(0, 5, bullet_text)
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    pdf.multi_cell(0, 5, "  " + line)
                else:
                    # Regular paragraph
                    pdf.multi_cell(0, 5, line)
            
            # Footer
            pdf.ln(10)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.ln(5)
            pdf.set_font('Helvetica', 'I', 8)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(0, 5, "HealthPulse USA - AI Healthcare Content Generator", align='C')
            
            return bytes(pdf.output())
            
        except Exception as e:
            print(f"PDF export error: {e}")
            return None


def get_download_filename(title: str, format: str) -> str:
    """Generate a clean filename for download"""
    # Clean the title for filename use
    clean_title = re.sub(r'[^\w\s-]', '', title)
    clean_title = re.sub(r'\s+', '_', clean_title)
    clean_title = clean_title[:50]  # Limit length
    
    timestamp = datetime.now().strftime("%Y%m%d")
    
    return f"HealthPulse_{clean_title}_{timestamp}.{format}"
